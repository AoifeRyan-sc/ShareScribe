# import datetime
import io
import base64
import openai
import pandas as pd
from .supported_languages import language_codes
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import dash_bootstrap_components as dbc
from dash import html, dcc

# from dotenv import load_dotenv
# load_dotenv()

def check_file(contents, filename):
    _, content_string = contents.split(',')
    file_extension = filename.split('.')[-1].lower()
    print("checking size")
    decoded = base64.b64decode(content_string)

    max_size = 25 * 1024 * 1024 # limit set by openai until I implement chunking 

    file_size = len(decoded)
    if file_size > max_size:
        return f'Error: {filename} exceeds the maximum file size of 25 MB.'
    elif file_extension not in ['wav', 'mp3', 'm4a']: # permitted file types
        return f'Error: Invalid file type. Please upload a .wav, .mp3, or .m4a file.'
    
    return decoded

def parse_contents(action, contents, transcribe_language = None):
    _, content_string = contents.split(',')
    print("parse_content running")
    print(f"language: {transcribe_language}")
    decoded = base64.b64decode(content_string)
    file_like = io.BytesIO(decoded)
    file_like.name = 'audio.m4a'

    api_key = os.getenv("OPENAI_API_KEY")
    
    client = openai.OpenAI(api_key = api_key)

    api_output = getattr(client.audio, action).create(
        model="whisper-1", 
        file=file_like,
        **({"language": transcribe_language} if action == "transcriptions" else {}),
        response_format = "srt"
    )

    return api_output

def translate_transcription(parsed_file, language_to, language_from, words):

    api_key = os.getenv("OPENAI_API_KEY")
    client = openai.OpenAI(api_key = api_key)

    pattern = r', ?'

    words_for_req = "Amazon, Diageo," + words if words is not None else "Amazon, Diageo"
    words_split = re.split(pattern, words_for_req) 
    print(words_split)
    words_quote_list = [f"'{word}'" for word in words_split]
    print(words_quote_list)
    words_string = ", ".join(words_quote_list)
    print(words_string)

    f_language_from = str.lower(language_codes[language_from])
    f_language_to = str.lower(language_codes[language_to])

    prompt = f"Translate the following SRT subtitle text from {f_language_from} to {f_language_to} as plain text. Don't translate proper nouns like {words_string}. Return only the translated SRT content without any formatting:\n\n {parsed_file}"
    # request = f"Below is an srt with {f_language_from} text. Translate it to an srt in {f_language_to}. Don't translate words that don't have a {f_language_to} translation, some examples are: {words_string}. \n\n {parsed_file}"
    
    print(f'request: {prompt}')

    print("-----------------")

    translated_srt = client.responses.create(
        model="gpt-4.1-mini",
        input=prompt
    )
    print("-----------------")

    print(f"translated_text: {translated_srt.output_text}")
    print(f"class translated_srt: {type(translated_srt)}")
    print(f"class translated_srt.output_text: {type(translated_srt.output_text)}")

    return translated_srt.output_text


def srt_to_docx(srt_string):

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    entries = re.split(r'\n\n', srt_string.strip())

    for entry in entries:
        lines = entry.split('\n')
        # if len(lines) >= 3:  # Ensure we have at least timestamp and text?
        timestamp = lines[1]
        text = ' '.join(lines[2:])

        p = doc.add_paragraph()

        timestamp_run = p.add_run(timestamp + " ")
        timestamp_run.font.color.rgb = RGBColor(192, 192, 192)  # Light grey

        # wrapped_text = textwrap.wrap(text, width=60)  # Adjust width as needed
        p.add_run(lines[2])
    
    byte_stream = io.BytesIO()
    doc.save(byte_stream)
    byte_stream.seek(0)  

        # for line in wrapped_text[1:]:
            # p.add_run('\n' + ' ' * 8 + line)

        # p.paragraph_format.space_after = Pt(6)  # Space after paragraph

    return byte_stream


def info_popover(icon_id, popover_text):
   
   pop = html.Span([
        # dbc.Label(className="fa fa-circle-info ms-2 d-inline-block", id="title-tooltip", html_for="page-title", style={"position": "relative", "top": "-1mm", 'color': '#1C7E75'}),
       dbc.Label(className="fa fa-circle-info ms-2 d-inline-block", id=icon_id, html_for="page-title", style={"position": "relative", "top": "-1mm", 'color': '#1C7E75'}),
        dbc.Tooltip(
            popover_text,
            # "Currently only accepting .wav, .mp3, and .m4a files, contact the Data Science team to add to compatible formats. If you have a video file you would like to transcribe, you can export as `audio only` from QuickTime (and other apps).",
            id = icon_id + "_hover",
            is_open=False,
            target = icon_id,
            placement="right"
        )
    ])
   
   return pop

def language_selection_dropdown(id, label):

    dd = dcc.Dropdown(
            options = language_codes,
            value = "",
            placeholder = label,
            # className="form-control",
            id = id,
            style = {'flex-grow': '1'}
        )
    
    return dd